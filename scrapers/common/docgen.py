#!/usr/bin/env python3
"""
Generate Word documents for each company in the Innovator repository from open rounds CSV data.

Reads the open_rounds CSV, locates each company's subfolder (Sector/CompanyName), and creates
a formatted .docx file with scraped company information.

Usage:
  python generate_company_docs.py [options]
  python generate_company_docs.py --csv ~/Downloads/Innovator/open_rounds_2026-02-22_23-57-59.csv
"""
import argparse
import csv
import glob
import logging
import os
import re
import sys

from docx import Document

# Reuse sanitization from scrapers.common.download
from scrapers.common.download import sanitize_filename, sanitize_foldername


def _sector(row: dict) -> str:
    """Derive sector from row: Sector > Product Development Stage > Regulatory Pathway > Round."""
    for k in ("Sector", "sector"):
        if k in row and (row.get(k) or "").strip():
            return (row.get(k) or "").strip()
    for col in ("Product Development Stage", "Regulatory Pathway", "Round"):
        if col in row:
            v = (row.get(col) or "").strip()
            if v:
                return v
    return "Open Rounds"


def _find_company_folder(output_dir: str, company: str, sector: str, create_folders: bool) -> str | None:
    """
    Return path to company folder, or None if not found.
    Tries computed path first, then searches all sector dirs for matching company folder.
    """
    sector_folder = sanitize_foldername(sector)
    company_folder = sanitize_filename(company)
    computed_path = os.path.join(output_dir, sector_folder, company_folder)

    if os.path.isdir(computed_path):
        return computed_path

    if create_folders:
        os.makedirs(computed_path, exist_ok=True)
        return computed_path

    # Search all sector subdirs for matching company folder
    for name in os.listdir(output_dir):
        sector_path = os.path.join(output_dir, name)
        if not os.path.isdir(sector_path):
            continue
        candidate = os.path.join(sector_path, company_folder)
        if os.path.isdir(candidate):
            return candidate

    return None


# Control chars and NULL bytes that are invalid in XML (python-docx uses XML)
_INVALID_XML_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")


def _sanitize_xml(s: str) -> str:
    """Remove NULL bytes and control characters invalid in XML."""
    if not s:
        return s
    return _INVALID_XML_RE.sub("", s)


def _get(row: dict, *keys: str) -> str:
    """Return first non-empty value from row for given keys (XML-safe)."""
    for k in keys:
        v = (row.get(k) or "").strip()
        if v:
            return _sanitize_xml(v)
    return ""


def _add_section(doc: Document, title: str, items: list[tuple[str, str]]) -> None:
    """Add a heading and labeled paragraphs. Skips empty values."""
    if not any(v for _k, v in items):
        return
    doc.add_heading(title, level=2)
    for label, value in items:
        if not value:
            continue
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)
    doc.add_paragraph()


def _add_paragraph_section(doc: Document, title: str, value: str) -> None:
    """Add a heading and a single paragraph."""
    if not value:
        return
    doc.add_heading(title, level=2)
    doc.add_paragraph(value)
    doc.add_paragraph()


def _add_team_section(doc: Document, team_members: str) -> None:
    """Add Team section with members split by semicolon."""
    if not team_members:
        return
    doc.add_heading("Team", level=2)
    for block in team_members.split(";"):
        block = block.strip()
        if not block:
            continue
        doc.add_paragraph(block)
    doc.add_paragraph()


def _add_links_section(doc: Document, company_url: str, video_url: str) -> None:
    """Add Links section."""
    if not company_url and not video_url:
        return
    doc.add_heading("Links", level=2)
    if company_url:
        doc.add_paragraph(f"Company: {company_url}")
    if video_url:
        doc.add_paragraph(f"Video: {video_url}")
    doc.add_paragraph()


def build_document(row: dict) -> Document:
    """Build a Word document from a CSV row."""
    doc = Document()

    company = _get(row, "Company Name")
    doc.add_heading(company, level=0)
    doc.add_paragraph()

    # Overview
    one_liner = _get(row, "One-liner")
    overview_items = [
        ("One-liner", one_liner),
        ("Website", _get(row, "Website")),
        ("Location", _get(row, "Location")),
        ("Year Founded", _get(row, "Year Founded")),
        ("Team Size", _get(row, "Team Size")),
        ("Current Runway", _get(row, "Current Runway")),
    ]
    _add_section(doc, "Overview", overview_items)

    # Deal
    deal_items = [
        ("Deal Summary", _get(row, "Deal Summary")),
        ("Urgency", _get(row, "Urgency")),
        ("Deal Type", _get(row, "Deal Type")),
        ("Round", _get(row, "Round")),
        ("Target Total", _get(row, "Target Total")),
        ("Open Amount", _get(row, "Open Amount")),
        ("Have Terms", _get(row, "Have Terms")),
    ]
    _add_section(doc, "Deal", deal_items)

    # Financing
    fin_items = [
        ("Total Equity To-date", _get(row, "Total Equity To-date")),
        ("Total Debt To-date", _get(row, "Total Debt To-date")),
        ("Total Non-Dilutive To-date", _get(row, "Total Non-Dilutive To-date")),
    ]
    _add_section(doc, "Financing", fin_items)

    # Product
    prod_items = [
        ("Primary Product Name", _get(row, "Primary Product Name")),
        ("Product Summary", _get(row, "Product Summary")),
    ]
    _add_section(doc, "Product", prod_items)

    # Stage & Regulatory
    reg_items = [
        ("Product Development Stage", _get(row, "Product Development Stage")),
        ("Regulatory Pathway", _get(row, "Regulatory Pathway")),
        ("US Regulatory Status", _get(row, "US Regulatory Status")),
        ("EU Regulatory Status", _get(row, "EU Regulatory Status")),
        ("Asia Regulatory Status", _get(row, "Asia Regulatory Status")),
    ]
    _add_section(doc, "Stage & Regulatory", reg_items)

    # Milestones
    milestone_items = [
        ("Milestones Completed", _get(row, "Milestones Completed")),
        ("Milestones Funded", _get(row, "Milestones Funded")),
        ("Milestones Open Round", _get(row, "Milestones Open Round")),
    ]
    _add_section(doc, "Milestones", milestone_items)

    # Team
    _add_team_section(doc, _get(row, "Team Members"))

    # Links
    _add_links_section(doc, _get(row, "Company URL"), _get(row, "Video URL"))

    return doc


def newest_open_rounds_csv(output_dir: str) -> str | None:
    """Return path to newest open_rounds_*.csv in output_dir, or None."""
    pattern = os.path.join(output_dir, "open_rounds_*.csv")
    candidates = glob.glob(pattern)
    if not candidates:
        return None
    return max(candidates, key=os.path.getmtime)


def main() -> None:
    default_output = os.path.join(os.path.expanduser("~"), "Downloads", "Innovator")
    default_csv = newest_open_rounds_csv(default_output) if os.path.isdir(default_output) else None

    parser = argparse.ArgumentParser(
        description="Generate Word docs per company from open rounds CSV into Innovator folder structure."
    )
    parser.add_argument(
        "--csv",
        default=default_csv,
        help="Path to open_rounds CSV (default: newest open_rounds_*.csv in output dir)",
    )
    parser.add_argument(
        "--output-dir",
        default=default_output,
        help="Innovator base directory (default: ~/Downloads/Innovator)",
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="List companies and target paths, no writes",
    )
    parser.add_argument(
        "--create-folders",
        action="store_true",
        help="Create sector/company folder if missing",
    )
    parser.add_argument(
        "--overwrite",
        action="store_true",
        help="Overwrite existing docs",
    )
    parser.add_argument(
        "-v", "--verbose",
        action="store_true",
        help="Verbose logging",
    )
    args = parser.parse_args()

    logging.basicConfig(
        level=logging.DEBUG if args.verbose else logging.INFO,
        format="%(message)s",
    )

    if not args.csv or not os.path.isfile(args.csv):
        logging.error("CSV file not found: %s", args.csv or "(none)")
        sys.exit(1)

    output_dir = os.path.abspath(args.output_dir)
    if not os.path.isdir(output_dir):
        if args.create_folders:
            os.makedirs(output_dir, exist_ok=True)
        else:
            logging.error("Output directory not found: %s", output_dir)
            sys.exit(1)

    with open(args.csv, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        fieldnames = list(reader.fieldnames or [])
        rows = list(reader)

    if not rows:
        logging.info("No rows in CSV.")
        return

    written = 0
    skipped_no_folder = 0
    skipped_exists = 0
    failed = 0

    for row in rows:
        company = (row.get("Company Name") or "").strip()
        if not company:
            continue

        sector = _sector(row)
        folder_path = _find_company_folder(output_dir, company, sector, args.create_folders)

        if not folder_path:
            skipped_no_folder += 1
            logging.debug("No folder for %s (sector=%s)", company, sector)
            continue

        doc_filename = sanitize_filename(company) + ".docx"
        doc_path = os.path.join(folder_path, doc_filename)

        if args.dry_run:
            logging.info("Would write: %s", doc_path)
            written += 1
            continue

        if not args.overwrite and os.path.isfile(doc_path):
            skipped_exists += 1
            continue

        try:
            doc = build_document(row)
            doc.save(doc_path)
            written += 1
            if args.verbose:
                logging.info("Wrote: %s", doc_path)
        except Exception as e:
            failed += 1
            logging.warning("Failed %s: %s", company, e)

    if args.dry_run:
        logging.info("Dry run: would write %d doc(s).", written)
    else:
        logging.info(
            "Done: %d written, %d skipped (existing), %d skipped (no folder), %d failed.",
            written, skipped_exists, skipped_no_folder, failed,
        )


if __name__ == "__main__":
    main()
