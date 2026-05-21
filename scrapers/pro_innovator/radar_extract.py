#!/usr/bin/env python3
"""
Radar Forum offline pipeline:

- Parse saved "Innovator Portal*.html/htm" (Save As -> Webpage, Complete)
- Write a merged radar_companies.csv
- Generate a media index CSV compatible with download_from_index.py
- Generate a simple per-company Word doc (.docx) from the extracted fields
- Emit helper scripts to download media

This is intentionally offline: it only needs the saved HTML snapshots.
"""

from __future__ import annotations

import argparse
import csv
import os
import re
from pathlib import Path

from docx import Document

from scrapers.pro_innovator.radar_portal_reader import (
    merge_records,
    parse_radar_html,
    write_csv,
)


MEDIA_SOURCES: list[tuple[str, str]] = [
    ("Pitch Deck", "Pitch Deck URL"),
    ("5 Minute Pitch Recording", "Pitch Recording URL"),
    ("Pitch Event Video", "Pitch Event Video URL"),
    ("Product Video", "Product Video URL"),
]


INDEX_HEADERS = [
    "Company Name",
    "Sector",
    "Link Label",
    "PDF URL",
    # keep a few useful fields for filtering/debugging (ignored by downloader)
    "Round",
    "Clinical Areas",
    "Country",
]


def _safe_name(s: str, max_len: int = 120) -> str:
    s = (s or "").strip()
    s = re.sub(r"[#%&*:<>?\"/\\\\|]+", "_", s)
    s = re.sub(r"\s+", " ", s).strip().strip(".")
    return (s[:max_len] if len(s) > max_len else s) or "Unknown"


def _pick_sector(rec: dict) -> str:
    for k in ("Clinical Areas", "Round"):
        v = (rec.get(k) or "").strip()
        if v:
            return v
    return "Radar"


def write_media_index(records: list[dict], output_path: Path) -> int:
    rows: list[dict] = []
    for rec in records:
        company = (rec.get("Company Name") or "").strip()
        if not company:
            continue
        sector = _pick_sector(rec)
        base = {
            "Company Name": company,
            "Sector": sector,
            "Round": (rec.get("Round") or "").strip(),
            "Clinical Areas": (rec.get("Clinical Areas") or "").strip(),
            "Country": (rec.get("Country") or "").strip(),
        }
        for label, url_field in MEDIA_SOURCES:
            url = (rec.get(url_field) or "").strip()
            if not url:
                continue
            rows.append(
                {
                    **base,
                    "Link Label": label,
                    "PDF URL": url,
                }
            )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", newline="", encoding="utf-8") as f:
        w = csv.DictWriter(f, fieldnames=INDEX_HEADERS, extrasaction="ignore")
        w.writeheader()
        w.writerows(rows)
    return len(rows)


def _add_kv_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in items:
        if not (k or "").strip():
            continue
        if not (v or "").strip():
            continue
        row = table.add_row().cells
        row[0].text = str(k).strip()
        row[1].text = str(v).strip()


def write_company_docs(records: list[dict], output_dir: Path) -> int:
    """
    Write one .docx per company into:
      output_dir/<Company Name>/<Company Name>.docx
    """
    count = 0
    for rec in records:
        company = (rec.get("Company Name") or "").strip()
        if not company:
            continue
        company_dir = output_dir / _safe_name(company)
        company_dir.mkdir(parents=True, exist_ok=True)
        doc = Document()
        doc.add_heading(company, level=0)

        highlights: list[tuple[str, str]] = []
        for k in (
            "Product Name (Grid)",
            "Product Name",
            "Short Description",
            "One-Line Description",
            "Product Abstract",
            "Product Description",
            "Product Development Stage",
            "On Market",
            "Year Founded",
            "Country",
            "Detail Country",
            "State",
            "City",
            "Clinical Areas",
            "Round",
            "Website",
        ):
            v = (rec.get(k) or "").strip()
            if v:
                highlights.append((k, v))

        if highlights:
            doc.add_heading("Details", level=2)
            _add_kv_table(doc, highlights)

        links: list[tuple[str, str]] = []
        for label, url_field in MEDIA_SOURCES:
            url = (rec.get(url_field) or "").strip()
            if url:
                links.append((label, url))
        if links:
            doc.add_heading("Links", level=2)
            _add_kv_table(doc, links)

        out = company_dir / f"{_safe_name(company)}.docx"
        doc.save(str(out))
        count += 1

    return count


def write_helper_scripts(output_dir: Path, index_csv: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)

    # macOS-friendly double-clickable script
    command_path = output_dir / "download_media.command"
    command_path.write_text(
        "\n".join(
            [
                "#!/bin/zsh",
                "set -euo pipefail",
                "",
                f"cd {output_dir.as_posix()!s}",
                "",
                "# Ensure we run from the repo root so imports work",
                f"REPO_ROOT={Path(__file__).resolve().parents[1].as_posix()!s}",
                "cd \"$REPO_ROOT/resi\"",
                "",
                "# Use a local venv so yt-dlp runs with modern SSL/Python",
                f"VENV_DIR=\"{(output_dir / '.venv_media').as_posix()}\"",
                "if [[ ! -x \"$VENV_DIR/bin/python\" ]]; then",
                "  echo \"Creating venv at: $VENV_DIR\"",
                "  python3 -m venv \"$VENV_DIR\"",
                "  \"$VENV_DIR/bin/python\" -m pip install --upgrade pip >/dev/null",
                "  \"$VENV_DIR/bin/python\" -m pip install yt-dlp >/dev/null",
                "fi",
                "",
                f"\"$VENV_DIR/bin/python\" download_from_index.py {index_csv.as_posix()!s} --output-dir {output_dir.as_posix()!s} --flat",
                "",
                "# Optional: convert per-company .docx -> PDFs (requires LibreOffice)",
                f"\"$VENV_DIR/bin/python\" convert_docx_to_pdf.py --docx-dir \"{output_dir.as_posix()}\" --out-dir \"{output_dir.as_posix()}\" --recursive || true",
                "",
                "echo \"\"",
                "echo \"Done. Output is organized by Sector/Company under:\"",
                f"echo \"  {output_dir.as_posix()!s}\"",
                "read -r \"?Press Enter to close...\"",
            ]
        )
        + "\n",
        encoding="utf-8",
    )
    os.chmod(command_path, 0o755)


def run(html_files: list[str], output_dir: Path) -> tuple[Path, Path]:
    all_records = []
    for p in html_files:
        records = parse_radar_html(p)
        all_records.append(records)
    merged = merge_records(all_records)

    radar_csv = output_dir / "radar_companies.csv"
    write_csv(merged, str(radar_csv))

    media_index = output_dir / "radar_media_index.csv"
    write_media_index(merged, media_index)

    write_company_docs(merged, output_dir)
    write_helper_scripts(output_dir, media_index)

    return radar_csv, media_index


def main() -> None:
    ap = argparse.ArgumentParser(
        description="Extract Radar HTML -> CSV, media index, and company docs.",
    )
    ap.add_argument(
        "html",
        nargs="+",
        help="Path(s) to saved Innovator Portal Radar HTML (.html/.htm).",
    )
    ap.add_argument(
        "--output-dir",
        "-o",
        required=True,
        type=Path,
        help="Output directory (e.g. ~/Downloads/virtual0416).",
    )
    args = ap.parse_args()

    out = args.output_dir.expanduser().resolve()
    out.mkdir(parents=True, exist_ok=True)
    run(args.html, out)


if __name__ == "__main__":
    main()

