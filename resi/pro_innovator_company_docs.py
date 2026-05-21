#!/usr/bin/env python3
"""
Generate one Word .docx "company brief" per company from a Pro Innovator companies CSV.

Writes into the existing downloader folder structure:
  OUTPUT_DIR/<Sector>/<Company>/<Company>_brief.docx

This is preferred over PDF briefs when formatting/wrapping must be robust.
"""

from __future__ import annotations

import argparse
import csv
import sys
from pathlib import Path


def _safe(s: str) -> str:
    s = (s or "").replace("\r", " ").replace("\n", " ").strip()
    return " ".join(s.split())


def _is_internal_grid_key(key: str) -> bool:
    return key.startswith("Grid: ")


def _company_folder(output_dir: Path, sector: str, company_name: str) -> Path:
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from download_from_index import sanitize_filename, sanitize_foldername  # noqa

    sector_dir = output_dir / sanitize_foldername(sector)
    return sector_dir / sanitize_filename(company_name)


def _brief_doc_path(folder: Path, company_name: str) -> Path:
    sys.path.insert(0, str(Path(__file__).resolve().parent))
    from download_from_index import sanitize_filename  # noqa

    safe = sanitize_filename(company_name, max_length=140)
    return folder / f"{safe}_brief.docx"


def _add_heading(doc, text: str) -> None:
    doc.add_heading(text, level=0)


def _add_section(doc, title: str, items: list[tuple[str, str]]) -> None:
    active = [(k, v) for k, v in items if (v or "").strip()]
    if not active:
        return
    doc.add_heading(title, level=2)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for k, v in active:
        row = table.add_row().cells
        row[0].text = str(k).strip()
        row[1].text = str(v).strip()


def _ordered_items(row: dict, include_grid: bool) -> list[tuple[str, str]]:
    keys = list(row.keys())
    preferred = [
        "One-liner",
        "One-Line Description",
        "Short Description",
        "Product Name",
        "Product Description",
        "Product abstract",
        "Product Abstract",
        "Development Stage",
        "Product development stage",
        "Product Development Stage",
        "Website",
        "Link to your organization's website",
        "Link to your organization's website URL",
        "Website URL",
        "Country",
        "Detail Country",
        "State",
        "City",
        "Year Founded",
        "What year was your organization founded?",
        "Is this product currently on the market?",
        "Round / Type",
        "Amount (Millions **USD**)",
        "Anticipated Close Date",
    ]

    seen = set()
    ordered: list[str] = []
    for k in preferred:
        if k in row and k not in seen:
            ordered.append(k)
            seen.add(k)

    # Put all media/link URL fields next.
    for k in keys:
        if k.endswith(" URL") and k not in seen:
            ordered.append(k)
            seen.add(k)

    # Everything else afterwards.
    for k in keys:
        if k in seen:
            continue
        if not include_grid and _is_internal_grid_key(k):
            continue
        if k == "Company Name":
            continue
        ordered.append(k)
        seen.add(k)

    items: list[tuple[str, str]] = []
    for k in ordered:
        v = _safe(row.get(k) or "")
        if v:
            items.append((k, v))
    return items


def build_company_doc(row: dict, include_grid: bool = False):
    try:
        from docx import Document  # type: ignore
    except Exception:
        raise SystemExit(
            "Missing dependency `python-docx`. Install with: python3 -m pip install --user python-docx"
        )

    doc = Document()
    company = _safe(row.get("Company Name") or "Unknown")
    _add_heading(doc, company)

    # High-signal sections first
    overview = [
        ("One-Line Description", _safe(row.get("One-Line Description") or row.get("Short Description") or "")),
        ("Website", _safe(row.get("Website URL") or row.get("Link to your organization's website URL") or row.get("Website") or "")),
        ("Country", _safe(row.get("Country") or row.get("Detail Country") or "")),
        ("State", _safe(row.get("State") or "")),
        ("City", _safe(row.get("City") or "")),
        ("Year Founded", _safe(row.get("Year Founded") or row.get("What year was your organization founded?") or "")),
    ]
    _add_section(doc, "Company Overview", overview)

    product = [
        ("Product Name", _safe(row.get("Product Name") or "")),
        ("Development Stage", _safe(row.get("Product Development Stage") or row.get("Product development stage") or row.get("Development Stage") or "")),
        ("On Market", _safe(row.get("Is this product currently on the market?") or "")),
    ]
    _add_section(doc, "Product", product)

    # Add all remaining fields as a catch-all table.
    _add_section(doc, "All Fields", _ordered_items(row, include_grid=include_grid))
    return doc


def run(companies_csv: Path, output_dir: Path, sector: str, include_grid: bool = False) -> int:
    with companies_csv.open(newline="", encoding="utf-8-sig") as f:
        rows = list(csv.DictReader(f))

    count = 0
    for row in rows:
        company = _safe(row.get("Company Name") or "")
        if not company:
            continue
        folder = _company_folder(output_dir, sector, company)
        folder.mkdir(parents=True, exist_ok=True)
        out_path = _brief_doc_path(folder, company)
        doc = build_company_doc(row, include_grid=include_grid)
        doc.save(str(out_path))
        count += 1
    return count


def main() -> None:
    ap = argparse.ArgumentParser(description="Write per-company Word briefs (.docx) from a Pro Innovator companies CSV.")
    ap.add_argument("--companies-csv", required=True, type=Path, help="Path to *_companies.csv")
    ap.add_argument("--output-dir", "-o", required=True, type=Path, help="Output dir used by downloader")
    ap.add_argument("--sector", required=True, help="Sector folder (e.g. APAC or MTI)")
    ap.add_argument("--include-grid", action="store_true", help="Include internal Grid:* fields")
    args = ap.parse_args()

    n = run(
        companies_csv=args.companies_csv.expanduser().resolve(),
        output_dir=args.output_dir.expanduser().resolve(),
        sector=args.sector,
        include_grid=args.include_grid,
    )
    print(f"Wrote {n} DOCX brief(s).")


if __name__ == "__main__":
    main()

