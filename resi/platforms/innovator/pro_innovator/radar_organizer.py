#!/usr/bin/env python3
"""
Organize Radar Forum data into per-company folders with Word docs and downloaded media.

Reads the radar_companies.csv produced by radar_portal_reader.py and creates:

  OUTPUT_DIR/
    radar_companies.csv                 # master CSV (copied)
    [Clinical Areas]/
      [CompanyName]/
        [CompanyName].docx              # summary doc
        [CompanyName]_pitch_deck.pptx   # or .pdf
        [CompanyName]_pitch_recording.mp4
        [CompanyName]_event_video.mp4   # if available
        [CompanyName]_product_video.mp4 # YouTube, if yt-dlp installed

Usage:
    python -m platforms.innovator.pro_innovator.radar_organizer [--csv PATH] [-o DIR]
    python -m platforms.innovator.pro_innovator.radar_organizer --docs-only
    python -m platforms.innovator.pro_innovator.radar_organizer --no-youtube
"""

import argparse
import csv
import logging
import os
import re
import sys
from pathlib import Path
from urllib.parse import urlparse, parse_qs, unquote

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[3]))
from download_from_index import (
    download_media,
    download_youtube,
    sanitize_filename,
    sanitize_foldername,
)
from urllib.parse import quote

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    datefmt="%H:%M:%S",
)
log = logging.getLogger(__name__)

DEFAULT_OUTPUT = os.path.join(os.path.expanduser("~"), "Downloads", "radar")
DEFAULT_CSV = os.path.join(DEFAULT_OUTPUT, "radar_companies.csv")

_INVALID_XML_RE = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]")


def _xml_safe(s: str) -> str:
    return _INVALID_XML_RE.sub("", s) if s else s


def _get(row: dict, *keys: str) -> str:
    for k in keys:
        v = (row.get(k) or "").strip()
        if v:
            return _xml_safe(v)
    return ""


# ---------------------------------------------------------------------------
# URL helpers
# ---------------------------------------------------------------------------

def _encode_url_spaces(url: str) -> str:
    """Percent-encode only literal spaces; leave everything else (including existing %-encoding) intact."""
    if not url:
        return url
    return url.replace(" ", "%20")


def extract_deck_download_url(viewer_url: str) -> tuple[str, str]:
    """
    Extract the actual downloadable file URL and extension from a viewer wrapper URL.
    Returns (download_url, extension).
    """
    if not viewer_url:
        return "", ""

    parsed = urlparse(viewer_url)

    if "officeapps.live.com" in parsed.netloc:
        src = parse_qs(parsed.query).get("src", [""])[0]
        if src:
            ext = ".pptx" if ".pptx" in src.lower() else ".pdf"
            return _encode_url_spaces(src), ext

    if "file-viewer" in parsed.path:
        url_param = parse_qs(parsed.query).get("url", [""])[0]
        if url_param:
            ext = ".pdf" if ".pdf" in url_param.lower() else ".pptx"
            return _encode_url_spaces(url_param), ext

    if "mti-innovator.s3" in viewer_url or "media.innovator.org" in viewer_url:
        ext = ".pptx" if ".pptx" in viewer_url.lower() else ".pdf"
        return _encode_url_spaces(viewer_url), ext

    return _encode_url_spaces(viewer_url), ".pdf"


# ---------------------------------------------------------------------------
# Word doc builder for Radar company data
# ---------------------------------------------------------------------------

def _add_section(doc: Document, title: str, items: list[tuple[str, str]]) -> None:
    active = [(l, v) for l, v in items if v]
    if not active:
        return
    doc.add_heading(title, level=2)
    for label, value in active:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(value)


def _add_long_text(doc: Document, title: str, text: str) -> None:
    if not text:
        return
    doc.add_heading(title, level=2)
    doc.add_paragraph(text)


def build_radar_doc(row: dict) -> Document:
    doc = Document()
    company = _get(row, "Company Name")
    product = _get(row, "Product Name", "Product Name (Grid)")

    doc.add_heading(company, level=0)
    if product:
        doc.add_paragraph(product).italic = True

    _add_section(doc, "Company Overview", [
        ("Short Description", _get(row, "Short Description", "One-Line Description")),
        ("Website", _get(row, "Website URL", "Website")),
        ("Country", _get(row, "Country", "Detail Country")),
        ("State", _get(row, "State")),
        ("City", _get(row, "City")),
        ("Year Founded", _get(row, "Year Founded")),
        ("Clinical Areas", _get(row, "Clinical Areas")),
        ("Funding Round", _get(row, "Round")),
    ])

    _add_section(doc, "Product Information", [
        ("Product Name", _get(row, "Product Name")),
        ("Development Stage", _get(row, "Product Development Stage")),
        ("On Market", _get(row, "On Market")),
    ])

    _add_long_text(doc, "One-Line Description", _get(row, "One-Line Description"))
    _add_long_text(doc, "Product Abstract", _get(row, "Product Abstract"))
    _add_long_text(doc, "Product Description", _get(row, "Product Description"))

    media_items = [
        ("Pitch Deck", _get(row, "Pitch Deck")),
        ("Pitch Recording", _get(row, "5 Minute Pitch Recording")),
        ("Product Video", _get(row, "Product Video URL")),
    ]
    password = _get(row, "Product Video Password")
    if password:
        media_items.append(("Video Password", password))
    _add_section(doc, "Media & Links", media_items)

    return doc


# ---------------------------------------------------------------------------
# Organizer
# ---------------------------------------------------------------------------

def organize(
    csv_path: str,
    output_dir: str,
    skip_existing: bool = True,
    docs_only: bool = False,
    no_youtube: bool = False,
) -> None:
    out = os.path.abspath(output_dir)
    os.makedirs(out, exist_ok=True)

    with open(csv_path, newline="", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        rows = list(reader)

    if not rows:
        log.info("No rows in CSV.")
        return

    log.info("Processing %d companies from %s", len(rows), csv_path)

    stats = {"docs": 0, "decks": 0, "recordings": 0, "event_vids": 0,
             "yt_vids": 0, "skipped": 0, "failed": 0}

    for row in rows:
        company = (row.get("Company Name") or "").strip()
        if not company:
            continue

        company_folder = sanitize_filename(company)
        company_path = os.path.join(out, company_folder)
        os.makedirs(company_path, exist_ok=True)

        # --- Word doc ---
        doc_path = os.path.join(company_path, company_folder + ".docx")
        if not skip_existing or not os.path.isfile(doc_path):
            try:
                doc = build_radar_doc(row)
                doc.save(doc_path)
                stats["docs"] += 1
            except Exception as e:
                log.warning("Doc failed for %s: %s", company, e)
                stats["failed"] += 1
        else:
            stats["skipped"] += 1

        if docs_only:
            continue

        # --- Pitch Deck ---
        deck_viewer_url = (row.get("Pitch Deck URL") or "").strip()
        if deck_viewer_url:
            deck_url, deck_ext = extract_deck_download_url(deck_viewer_url)
            if deck_url:
                deck_path = os.path.join(
                    company_path, company_folder + "_pitch_deck" + deck_ext
                )
                if not skip_existing or not os.path.isfile(deck_path):
                    log.info("Deck: %s", company)
                    if download_media(deck_url, deck_path):
                        stats["decks"] += 1
                    else:
                        stats["failed"] += 1
                else:
                    stats["skipped"] += 1

        # --- Pitch Recording ---
        rec_url = (row.get("Pitch Recording URL") or "").strip()
        if rec_url:
            rec_path = os.path.join(
                company_path, company_folder + "_pitch_recording.mp4"
            )
            if not skip_existing or not os.path.isfile(rec_path):
                log.info("Recording: %s", company)
                if download_media(_encode_url_spaces(rec_url), rec_path):
                    stats["recordings"] += 1
                else:
                    stats["failed"] += 1
            else:
                stats["skipped"] += 1

        # --- Pitch Event Video ---
        event_url = (row.get("Pitch Event Video URL") or "").strip()
        if event_url and "media.innovator.org" in event_url:
            event_path = os.path.join(
                company_path, company_folder + "_event_video.mp4"
            )
            if not skip_existing or not os.path.isfile(event_path):
                log.info("Event video: %s", company)
                if download_media(_encode_url_spaces(event_url), event_path):
                    stats["event_vids"] += 1
                else:
                    stats["failed"] += 1
            else:
                stats["skipped"] += 1

        # --- Product Video (YouTube) ---
        if not no_youtube:
            prod_vid_url = (row.get("Product Video URL") or "").strip()
            if prod_vid_url and ("youtube.com" in prod_vid_url or "youtu.be" in prod_vid_url):
                yt_path = os.path.join(
                    company_path, company_folder + "_product_video.mp4"
                )
                if not skip_existing or not os.path.isfile(yt_path):
                    log.info("YouTube: %s", company)
                    if download_youtube(prod_vid_url, yt_path):
                        stats["yt_vids"] += 1
                    else:
                        stats["failed"] += 1
                else:
                    stats["skipped"] += 1

    log.info(
        "Done: %d docs, %d decks, %d recordings, %d event videos, %d YouTube videos, "
        "%d skipped, %d failed.",
        stats["docs"], stats["decks"], stats["recordings"], stats["event_vids"],
        stats["yt_vids"], stats["skipped"], stats["failed"],
    )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Organize Radar companies into per-company folders with docs and media."
    )
    parser.add_argument(
        "--csv", default=DEFAULT_CSV,
        help="Path to radar_companies.csv (default: ~/Downloads/radar/radar_companies.csv)",
    )
    parser.add_argument(
        "-o", "--output-dir", default=DEFAULT_OUTPUT,
        help="Output directory (default: ~/Downloads/radar)",
    )
    parser.add_argument(
        "--docs-only", action="store_true",
        help="Only generate Word docs, skip media downloads",
    )
    parser.add_argument(
        "--no-youtube", action="store_true",
        help="Skip YouTube product video downloads",
    )
    parser.add_argument(
        "--no-skip-existing", action="store_true",
        help="Re-download / overwrite existing files",
    )
    args = parser.parse_args()

    organize(
        csv_path=args.csv,
        output_dir=args.output_dir,
        skip_existing=not args.no_skip_existing,
        docs_only=args.docs_only,
        no_youtube=args.no_youtube,
    )


if __name__ == "__main__":
    main()
