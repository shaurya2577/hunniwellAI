#!/usr/bin/env python3
"""
PitchBook to Word Converter - Web UI
Run this script, then open http://127.0.0.1:5000 in your browser.
Drag an HTML file into the drop zone, choose where to save, then convert.
"""

import io
import re
import sys
from pathlib import Path
from html import unescape

# Web server - use Flask if available, else built-in http.server with a simple form
try:
    from flask import Flask, request, send_file
    FLASK_AVAILABLE = True
except ImportError:
    FLASK_AVAILABLE = False


# --- Conversion logic ---

def strip_html(text: str) -> str:
    text = re.sub(r"<[^>]+>", " ", text)
    text = re.sub(r"\s+", " ", text).strip()
    return unescape(text)


def extract_entity_name(html: str) -> str:
    """Extract name from title. Handles Company Profile, Investor Profile, and (EXCH:TICKER) format."""
    m = re.search(r"<title>([^<]+)\s*-\s*(?:Company|Investor)\s+Profile</title>", html, re.I)
    if m:
        return strip_html(m.group(1))
    return "Profile"


def extract_ai_profile_summary(html: str) -> list[str]:
    """Extract AI Profile Summary bullet points (narrative overview)."""
    bullets = []
    # Find the ul with summary-list that follows "AI Profile Summary" heading
    idx = html.find("AI Profile Summary")
    if idx < 0:
        return bullets
    # Locate summary-list ul (may have hashed class: summary-list_bulleted_ea76bcd1)
    ul_match = re.search(r'<ul[^>]*summary-list[^>]*>([\s\S]*?)</ul>', html[idx:])
    if not ul_match:
        return bullets
    for li in re.findall(r'<li[^>]*>([\s\S]*?)</li>', ul_match.group(1)):
        text = strip_html(li)
        if text and len(text) > 10:
            bullets.append(text)
    return bullets


def extract_sections(html: str) -> list[tuple[str, str]]:
    section_headers = list(
        re.finditer(r'typo-heading-level-3">([^<]+)</div></h3>', html)
    )
    if not section_headers:
        section_headers = list(
            re.finditer(r'font-weight-semi-bold typo-heading-level-3">([^<]+)</div>', html)
        )
    section_headers = [m for m in section_headers if m.group(1).strip()]
    sections = []
    for i, match in enumerate(section_headers):
        title = unescape(strip_html(match.group(1)))
        start = match.end()
        end = section_headers[i + 1].start() if i + 1 < len(section_headers) else len(html)
        sections.append((title, html[start:end]))
    return sections


def extract_tables_from_html(html: str) -> list[list[list[str]]]:
    tables = []
    for block in re.findall(
        r'<table[^>]*class="[^"]*table[^"]*"[^>]*>(.*?)</table>',
        html, re.DOTALL,
    ):
        rows = []
        for tr in re.findall(r"<tr[^>]*>(.*?)</tr>", block, re.DOTALL):
            cells = [strip_html(td) for td in re.findall(r"<t[dh][^>]*>(.*?)</t[dh]>", tr, re.DOTALL)]
            if cells:
                rows.append(cells)
        if rows:
            tables.append(rows)
    return tables


def _extract_highlight_tiles(html: str) -> list[tuple[str, str]]:
    """Extract highlight-tile title/value pairs (Highlights section).
    Handles hashed class names (e.g. highlight-tile__title_c47f88eb) and nested value structure."""
    titles = re.findall(r'highlight-tile__title[^>]*>([^<]*)', html)
    # Value can be direct (<h3>7</h3>) or nested (<h3><div><div class="group__i">$3.77Tn</div>...)
    values = []
    for m in re.finditer(r'highlight-tile__value[^>]*>([\s\S]*?)</h3>', html):
        inner = m.group(1)
        # Prefer group__i content (main number), else strip all tags
        gi = re.search(r'<div[^>]*group__i[^>]*>([^<]+)</div>', inner)
        if gi:
            values.append(strip_html(gi.group(1)))
        else:
            values.append(strip_html(inner))
    pairs = []
    for t, v in zip(titles, values):
        t, v = strip_html(t), (v or "").strip()
        if t or v:
            pairs.append((t or "—", v or "—"))
    return pairs


def _extract_table_list(html: str) -> list[list[str]]:
    """Extract table-list div structure (label/value rows).
    Rows have table-list__cell for label (or table-list__cell_label) and value."""
    rows = []
    # Find each row: content from table-list__row until next table-list__row
    pos = 0
    while True:
        m = re.search(r'<div[^>]*table-list__row[^>]*>', html[pos:])
        if not m:
            break
        row_start = pos + m.end()
        m2 = re.search(r'<div[^>]*table-list__row[^>]*>', html[row_start:])
        row_end = row_start + m2.start() if m2 else len(html)
        row_html = html[row_start:row_end]
        # Extract cells - label often in table-list__cell_label, value in sibling table-list__cell
        cells = []
        for cell_m in re.finditer(r'<div[^>]*table-list__cell[^>]*>([\s\S]*?)</div>', row_html):
            raw = cell_m.group(1)
            # Skip if this looks like an inner wrapper (e.g. span), take the leaf text
            text = strip_html(raw)
            if text:
                cells.append(text)
        if len(cells) >= 2 and any(c.strip() for c in cells):
            rows.append(cells[:2])
        pos = row_end
    return rows


def _extract_contact_info(html: str) -> list[list[str]]:
    """Extract Contact Information section (Primary Contact, Primary Office, etc.).
    Structure: label in ellipsis span, content in following ul/li."""
    rows = []
    for m in re.finditer(
        r'<span[^>]*ellipsis[^>]*>([^<]+)</span>[\s\S]*?<ul[^>]*>([\s\S]*?)</ul>',
        html, re.DOTALL,
    ):
        label = strip_html(m.group(1))
        items = re.findall(r'<li[^>]*>([\s\S]*?)</li>', m.group(2))
        text = "\n".join(strip_html(it) for it in items if strip_html(it))
        if label and text and len(text) > 2:
            rows.append([label, text])
    return rows


def _extract_data_points(html: str) -> list[tuple[str, str]]:
    pairs = []
    for block in re.findall(r'class="[^"]*data-point[^"]*"[^>]*>(.*?)</div>\s*</div>', html, re.DOTALL):
        label = re.search(r'data-point__label[^>]*>([^<]+)', block)
        val = re.search(r'data-point__(?:amount|value)[^>]*>([^<]*)', block)
        if label and val:
            pairs.append((strip_html(label.group(1)), strip_html(val.group(1))))
    return pairs


def _build_section_content(section_title: str, section_html: str) -> tuple[list[list[list[str]]], list[str]]:
    tables_data = extract_tables_from_html(section_html)
    paragraphs = []

    # Highlights: prefer highlight tiles
    if "Highlights" in section_title:
        tiles = _extract_highlight_tiles(section_html)
        if tiles:
            tables_data.insert(0, [["Metric", "Value"]] + [[k, v] for k, v in tiles])

    # Table-list (div-based key-value rows)
    table_list_rows = _extract_table_list(section_html)
    if table_list_rows:
        tables_data.append(table_list_rows)

    # Contact Information: label + ul/li structure (Primary Contact, Primary Office)
    if "Contact" in section_title:
        contact_rows = _extract_contact_info(section_html)
        if contact_rows:
            tables_data.append(contact_rows)

    # Data points
    data_points = _extract_data_points(section_html)
    if data_points and not tables_data:
        tables_data = [[["Metric", "Value"]]] + [[k, v] for k, v in data_points]

    if not tables_data:
        clean = re.sub(r"<script[^>]*>.*?</script>", "", section_html, flags=re.DOTALL)
        clean = re.sub(r"<style[^>]*>.*?</style>", "", clean, flags=re.DOTALL)
        for block in re.findall(r"<p[^>]*>(.*?)</p>", clean, re.DOTALL):
            text = strip_html(block)
            if len(text) > 20:
                paragraphs.append(text)
        for block in re.findall(r'class="[^"]*description[^"]*"[^>]*>(.*?)</div>', clean, re.DOTALL):
            text = strip_html(block)
            if len(text) > 30:
                paragraphs.append(text)
    return tables_data, paragraphs


def convert_to_docx_bytes(html_content: str) -> tuple[bytes, str, str]:
    """
    Convert HTML to Word document. Returns (file_bytes, filename, content_type).
    """
    entity_name = extract_entity_name(html_content)
    sections = extract_sections(html_content)

    if not sections:
        alt = re.findall(r'<div class="font-weight-semi-bold typo-heading-level-3">([^<]+)</div>', html_content)
        sections = [(unescape(h.strip()), "") for h in alt] if alt else [("Profile Data", html_content)]

    if not sections:
        raise ValueError("No structured data found in HTML")

    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading(entity_name, level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()

        # AI Profile Summary (if present)
        ai_bullets = extract_ai_profile_summary(html_content)
        if ai_bullets:
            doc.add_heading("AI Profile Summary", level=2)
            for bullet in ai_bullets:
                doc.add_paragraph(bullet[:32767], style="List Bullet")
            doc.add_paragraph()

        for section_title, section_html in sections:
            if not section_title.strip():
                continue
            doc.add_heading(section_title, level=2)
            tables_data, paragraphs = _build_section_content(section_title, section_html)
            for table_data in tables_data:
                if not table_data:
                    continue
                num_cols = max(len(row) for row in table_data)
                wtable = doc.add_table(rows=len(table_data), cols=num_cols)
                wtable.style = "Table Grid"
                for ri, row in enumerate(table_data):
                    for ci, cell_text in enumerate(row):
                        if ci < num_cols:
                            wtable.rows[ri].cells[ci].text = cell_text[:32767]
                doc.add_paragraph()  # spacing between tables
            for p in paragraphs:
                doc.add_paragraph(p)
            doc.add_paragraph()

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        safe_name = re.sub(r'[^\w\s\-\(\)]', '', entity_name).strip().replace(' ', '_')[:80] or "Profile"
        return buf.getvalue(), f"{safe_name}.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    except ImportError:
        pass

    # Fallback: HTML
    def _esc(s):
        return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;").replace('"', "&quot;")

    parts = [
        '<!DOCTYPE html><html xmlns:o="urn:schemas-microsoft-com:office:office" xmlns:w="urn:schemas-microsoft-com:office:word">',
        f"<head><meta charset='UTF-8'><title>{_esc(entity_name)}</title></head><body>",
        f"<h1 align='center'>{_esc(entity_name)}</h1><br/>",
    ]
    ai_bullets = extract_ai_profile_summary(html_content)
    if ai_bullets:
        parts.append("<h2>AI Profile Summary</h2><ul>")
        for bullet in ai_bullets:
            parts.append(f"<li>{_esc(bullet)}</li>")
        parts.append("</ul><br/>")
    for section_title, section_html in sections:
        if not section_title.strip():
            continue
        parts.append(f"<h2>{_esc(section_title)}</h2>")
        tables_data, paragraphs = _build_section_content(section_title, section_html)
        for table_data in tables_data:
            if not table_data:
                continue
            parts.append("<table border='1' cellpadding='4' cellspacing='0' style='border-collapse:collapse'>")
            for ri, row in enumerate(table_data):
                tag = "th" if ri == 0 else "td"
                parts.append("<tr>" + "".join(f"<{tag}>{_esc(c)}</{tag}>" for c in row) + "</tr>")
            parts.append("</table><br/>")
        for p in paragraphs:
            parts.append(f"<p>{_esc(p)}</p>")
        parts.append("<br/>")
    parts.append("</body></html>")

    safe_name = re.sub(r'[^\w\s\-\(\)]', '', entity_name).strip().replace(' ', '_')[:80] or "Profile"
    return "\n".join(parts).encode("utf-8"), f"{safe_name}.doc", "application/msword"


# --- Web UI ---

HTML_PAGE = """
<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>PitchBook → Word Converter</title>
    <style>
        * { box-sizing: border-box; }
        body { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', Roboto, sans-serif; margin: 0; padding: 24px; max-width: 560px; margin: 0 auto; }
        h1 { font-size: 1.5rem; margin-bottom: 8px; }
        .drop-zone { border: 2px dashed #999; border-radius: 12px; padding: 32px; text-align: center; cursor: pointer; transition: all 0.2s; background: #fafafa; }
        .drop-zone:hover, .drop-zone.dragover { border-color: #1d5080; background: #e8f4f8; }
        .drop-zone p { margin: 0; color: #666; }
        .drop-zone .file-name { margin-top: 8px; font-weight: 600; color: #1d5080; word-break: break-all; }
        .btn { display: inline-block; padding: 10px 20px; background: #1d5080; color: white; border: none; border-radius: 8px; cursor: pointer; font-size: 1rem; margin-top: 12px; }
        .btn:hover { background: #26649e; }
        .btn:disabled { background: #ccc; cursor: not-allowed; }
        .hidden { display: none; }
        #fileInput { display: none; }
        .log { margin-top: 16px; padding: 12px; background: #f5f5f5; border-radius: 8px; font-size: 0.9rem; min-height: 60px; }
        .log.success { color: #0a6; }
        .log.error { color: #c00; }
    </style>
</head>
<body>
    <h1>PitchBook → Word Converter</h1>
    <p>Drag your saved PitchBook HTML file here (or click to browse). Put the HTML file and its <code>_files</code> folder in the same place when saving from PitchBook—we read the HTML content to generate the Word doc.</p>

    <div id="dropZone" class="drop-zone">
        <p>Drop HTML file here</p>
        <p>or click to browse</p>
        <p class="file-name" id="fileName"></p>
    </div>
    <input type="file" id="fileInput" accept=".html,.htm">

    <p style="margin-top: 16px;">
        <button id="convertBtn" class="btn" disabled>Convert &amp; Download</button>
    </p>
    <p style="font-size:0.85rem;color:#666;margin-top:4px;">You'll choose where to save when the download dialog appears.</p>

    <div id="log" class="log"></div>

    <script>
        const dropZone = document.getElementById('dropZone');
        const fileInput = document.getElementById('fileInput');
        const fileName = document.getElementById('fileName');
        const convertBtn = document.getElementById('convertBtn');
        const log = document.getElementById('log');

        let selectedFile = null;

        function setFile(file) {
            selectedFile = file;
            fileName.textContent = file ? file.name : '';
            convertBtn.disabled = !file;
        }

        function logMsg(msg, type) {
            log.textContent = msg;
            log.className = 'log ' + (type || '');
        }

        dropZone.addEventListener('click', () => fileInput.click());

        dropZone.addEventListener('dragover', (e) => {
            e.preventDefault();
            dropZone.classList.add('dragover');
        });
        dropZone.addEventListener('dragleave', () => dropZone.classList.remove('dragover'));

        dropZone.addEventListener('drop', (e) => {
            e.preventDefault();
            dropZone.classList.remove('dragover');
            const items = e.dataTransfer.items;
            for (const item of items) {
                const file = item.getAsFile();
                if (file && file.name.match(/\\.html?$/i)) {
                    setFile(file);
                    return;
                }
            }
        });

        fileInput.addEventListener('change', () => {
            const f = fileInput.files[0];
            if (f) setFile(f);
        });

        convertBtn.addEventListener('click', async () => {
            if (!selectedFile) return;
            convertBtn.disabled = true;
            logMsg('Converting...');
            try {
                const formData = new FormData();
                formData.append('file', selectedFile);
                const res = await fetch('/convert', { method: 'POST', body: formData });
                const blob = await res.blob();
                if (!res.ok) {
                    const txt = await blob.text();
                    throw new Error(txt || res.statusText);
                }
                const cd = res.headers.get('Content-Disposition');
                let name = 'Company.docx';
                if (cd) {
                    const m = cd.match(/filename="?([^";\\n]+)"?/);
                    if (m) name = m[1];
                }
                const a = document.createElement('a');
                a.href = URL.createObjectURL(blob);
                a.download = name;
                a.click();
                URL.revokeObjectURL(a.href);
                logMsg('Download started: ' + name, 'success');
            } catch (err) {
                logMsg('Error: ' + err.message, 'error');
            }
            convertBtn.disabled = false;
        });
    </script>
</body>
</html>
"""


def create_app():
    app = Flask(__name__)

    @app.route("/")
    def index():
        return HTML_PAGE

    @app.route("/convert", methods=["POST"])
    def convert():
        if "file" not in request.files:
            return "No file uploaded", 400
        f = request.files["file"]
        if not f or not f.filename:
            return "No file selected", 400
        try:
            html_content = f.read().decode("utf-8", errors="replace")
            data, filename, ctype = convert_to_docx_bytes(html_content)
            return send_file(
                io.BytesIO(data),
                mimetype=ctype,
                as_attachment=True,
                download_name=filename,
            )
        except ValueError as e:
            return str(e), 400
        except Exception as e:
            return str(e), 500

    return app


def main():
    if not FLASK_AVAILABLE:
        print("Flask is required. Install with: pip install flask")
        sys.exit(1)

    import webbrowser
    from threading import Timer

    app = create_app()

    def open_browser():
        webbrowser.open("http://127.0.0.1:5000")

    Timer(1.0, open_browser).start()
    print("PitchBook Converter UI running at http://127.0.0.1:5000")
    print("Your browser should open automatically.")
    app.run(host="127.0.0.1", port=5000, debug=False, use_reloader=False)


if __name__ == "__main__":
    main()
